import barcode
from barcode.codex import Code128
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import for centering text
import datetime
import os
import time  # Import time module for sleep function
from pymongo import MongoClient

# Connect to MongoDB
client = MongoClient('mongodb://127.0.0.1:27017/')
db = client['CSVC']
device_collection = db['devices']

# Fetch all devices from the collection
devices = device_collection.find()

# Construct the items list
items = []
for idx, device in enumerate(devices, start=1):
    item = {
        'barcode': device['serialNumber'],
        'number': f'{idx:03}',
        'description': device['name']
    }
    items.append(item)

print(items)

# Create a Word document
doc = Document()
doc.add_heading('Barcodes', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Current date and time for the filename
now = datetime.datetime.now()
filename = f"Barcode-{now.strftime('%d%m%Y-%H%M')}.docx"

# Create a temporary folder to store barcode images
temp_folder = 'temp'
os.makedirs(temp_folder, exist_ok=True)

barcode_filenames = []  # List to store generated barcode filenames

# Generate and save barcode images
for item in items:
    # Generate barcode with number as the item code
    ean = Code128(f"{item['barcode']} - {item['number']}", writer=ImageWriter())
    
    # Save barcode as an image file in the temp folder
    barcode_filename = f"{item['barcode']}_{item['number']}"
    barcode_path = os.path.join(temp_folder, barcode_filename)
    ean.save(barcode_path)
    barcode_filenames.append(barcode_path)  # Store filename for deletion later

    # Add barcode information to the document
    doc.add_heading(f"Tên thiết bị: {item['description']}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph(f"Barcode: {item['barcode']} - {item['number']}")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(barcode_path + '.png', width=Inches(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")  # Add a blank paragraph for spacing

# Save the Word document
doc.save(filename)
print(f"Barcodes saved to {filename}")

# Delete the barcode image files after saving the document
for barcode_filename in barcode_filenames:
    # Check if the file exists before trying to remove it
    path = barcode_filename +  '.png'
    if os.path.exists(path):
        os.remove(path)
    else:
        print(f"File '{barcode_filename}' does not exist.")

# Remove the temporary folder
try:
    os.rmdir(temp_folder)
    print(f"Temporary folder '{temp_folder}' deleted.")
except OSError as e:
    print(f"Error: {temp_folder} : {e.strerror}")

print("Barcode images deleted.")
