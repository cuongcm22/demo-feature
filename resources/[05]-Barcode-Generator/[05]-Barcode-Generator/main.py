# import os
# import shutil
# import barcode
# from barcode.codex import Code128
# from barcode.writer import ImageWriter
# from docx import Document
# from docx.shared import Inches
# import datetime

# # Array of items
# items = ['dog', 'cat']

# # items = ['123456789', '987654321']

# # Create a Word document
# doc = Document()
# doc.add_heading('Barcodes', 0)

# # Current date and time for the filename
# now = datetime.datetime.now()
# filename = f"Barcode-{now.strftime('%d%m%Y-%H%M')}.docx"

# # Create a temporary folder to store barcode images
# temp_folder = 'temp'
# os.makedirs(temp_folder, exist_ok=True)

# barcode_filenames = []  # List to store generated barcode filenames

# # Generate and save barcode images
# for item in items:
#     # Generate barcode
#     ean = Code128(item, writer=ImageWriter())

#     # Save barcode as an image file in the temp folder
#     barcode_filename = f"{item}"
#     barcode_path = os.path.join(temp_folder, barcode_filename)
#     ean.save(barcode_path)

#     barcode_filenames.append(barcode_path)  # Store filename for deletion later

# # Add the barcode images to the Word document
# for item, barcode_filename in zip(items, barcode_filenames):
#     doc.add_heading(item, level=1)
#     doc.add_picture(barcode_filename + ".png", width=Inches(2))

# # Save the Word document
# doc.save(filename)
# print(f"Barcodes saved to {filename}")

# # Delete the barcode image files after saving the document
# for barcode_filename in barcode_filenames:
#     os.remove(barcode_filename + ".png")

# # Remove the temporary folder
# shutil.rmtree(temp_folder)
# print(f"Temporary folder '{temp_folder}' deleted.")

# print("Barcode images deleted.")


import barcode
from barcode.codex import Code128
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches
import datetime
import os
import time  # Import time module for sleep function

# Array of items with barcode, number, and text description
items = [
    {'barcode': '123456789', 'number': '001', 'description': 'Iphone 12'},
    {'barcode': '987654321', 'number': '002', 'description': 'Tivi LG'}
]

# Create a Word document
doc = Document()
doc.add_heading('Barcodes', 0)

# Current date and time for the filename
now = datetime.datetime.now()
filename = f"Barcode-{now.strftime('%d%m%Y-%H%M')}.docx"

# Create a temporary folder to store barcode images
temp_folder = 'temp'
os.makedirs(temp_folder, exist_ok=True)

barcode_filenames = []  # List to store generated barcode filenames

# Generate and save barcode images
for item in items:
    # Generate barcode with number as the item code
    ean = Code128(f"{item['barcode']} - {item['number']}", writer=ImageWriter())
    
    # Save barcode as an image file in the temp folder
    barcode_filename = f"{item['barcode']}_{item['number']}"
    barcode_path = os.path.join(temp_folder, barcode_filename)
    ean.save(barcode_path)
    barcode_filenames.append(barcode_path)  # Store filename for deletion later

    # Add barcode information to the document
    doc.add_heading(f"Item: {item['description']}", level=1)
    doc.add_paragraph(f"Barcode: {item['barcode']} - {item['number']}")
    print('barcode_path: ', barcode_path + '.png')
    doc.add_picture(barcode_path + '.png', width=Inches(2))
    doc.add_paragraph("")  # Add a blank paragraph for spacing

# Save the Word document
doc.save(filename)
print(f"Barcodes saved to {filename}")

# Delete the barcode image files after saving the document
for barcode_filename in barcode_filenames:
    # Check if the file exists before trying to remove it
    path = barcode_filename +  '.png'
    if os.path.exists(path):
        os.remove(path)
    else:
        print(f"File '{barcode_filename}' does not exist.")

# Remove the temporary folder
try:
    os.rmdir(temp_folder)
    print(f"Temporary folder '{temp_folder}' deleted.")
except OSError as e:
    print(f"Error: {temp_folder} : {e.strerror}")

print("Barcode images deleted.")
